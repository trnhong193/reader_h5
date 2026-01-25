#!/usr/bin/env python3
"""
Script đơn giản để chuyển đổi Markdown sang Word và PDF
Sử dụng pypandoc nếu có, nếu không thì dùng python-docx và markdown
"""

import os
import sys

def convert_with_pypandoc(md_file, word_file, pdf_file):
    """Chuyển đổi bằng pypandoc"""
    try:
        import pypandoc
        
        print("Đang sử dụng pypandoc...")
        
        # Chuyển sang Word
        print(f"Đang chuyển sang Word: {word_file}")
        pypandoc.convert_file(md_file, 'docx', outputfile=word_file, 
                             extra_args=['--standalone'])
        print(f"✓ Đã tạo file Word: {word_file}")
        
        # Chuyển sang PDF
        print(f"Đang chuyển sang PDF: {pdf_file}")
        pypandoc.convert_file(md_file, 'pdf', outputfile=pdf_file,
                             extra_args=['--pdf-engine=wkhtmltopdf', '--standalone'])
        print(f"✓ Đã tạo file PDF: {pdf_file}")
        
        return True
    except ImportError:
        return False
    except Exception as e:
        print(f"Lỗi với pypandoc: {e}")
        return False

def convert_with_python_docx(md_file, word_file):
    """Chuyển đổi sang Word bằng python-docx"""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import re
        
        print("Đang sử dụng python-docx...")
        
        doc = Document()
        
        # Đọc file markdown
        with open(md_file, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        
        in_code_block = False
        code_block_content = []
        code_language = ''
        in_table = False
        table_data = []
        
        i = 0
        while i < len(lines):
            line = lines[i].rstrip('\n')
            
            # Xử lý code blocks
            if line.strip().startswith('```'):
                if in_code_block:
                    # Kết thúc code block
                    if code_block_content:
                        para = doc.add_paragraph()
                        run = para.add_run('\n'.join(code_block_content))
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0, 0, 128)
                    code_block_content = []
                    in_code_block = False
                else:
                    in_code_block = True
                    code_language = line.strip()[3:].strip()
                i += 1
                continue
            
            if in_code_block:
                code_block_content.append(line)
                i += 1
                continue
            
            # Xử lý headers
            if line.startswith('# '):
                doc.add_heading(line[2:].strip(), level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:].strip(), level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:].strip(), level=3)
            elif line.startswith('#### '):
                doc.add_heading(line[5:].strip(), level=4)
            elif line.startswith('##### '):
                doc.add_heading(line[6:].strip(), level=5)
            elif line.startswith('###### '):
                doc.add_heading(line[7:].strip(), level=6)
            # Xử lý horizontal rules
            elif line.strip() == '---' or line.strip() == '***':
                para = doc.add_paragraph()
                para.add_run('_' * 50)
            # Xử lý tables
            elif '|' in line and line.strip().startswith('|'):
                # Bỏ qua separator line
                if not re.match(r'^\|[\s\-\|:]+\|$', line.strip()):
                    cells = [cell.strip() for cell in line.split('|')[1:-1]]
                    if not in_table:
                        # Bắt đầu table mới
                        table = doc.add_table(rows=1, cols=len(cells))
                        table.style = 'Light Grid Accent 1'
                        row = table.rows[0]
                        for j, cell_text in enumerate(cells):
                            cell = row.cells[j]
                            cell.text = cell_text
                            # Header row
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.bold = True
                        in_table = True
                    else:
                        # Thêm row mới
                        row = table.add_row()
                        for j, cell_text in enumerate(cells):
                            row.cells[j].text = cell_text
                else:
                    in_table = False
            # Xử lý danh sách
            elif line.strip().startswith('- '):
                doc.add_paragraph(line[2:].strip(), style='List Bullet')
                in_table = False
            elif re.match(r'^\d+\.\s', line.strip()):
                doc.add_paragraph(re.sub(r'^\d+\.\s', '', line.strip()), style='List Number')
                in_table = False
            # Xử lý code inline
            elif '`' in line:
                para = doc.add_paragraph()
                parts = re.split(r'(`[^`]+`)', line)
                for part in parts:
                    if part.startswith('`') and part.endswith('`'):
                        run = para.add_run(part[1:-1])
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
                    else:
                        para.add_run(part)
                in_table = False
            # Xử lý paragraph thông thường
            elif line.strip():
                doc.add_paragraph(line.strip())
                in_table = False
            # Dòng trống
            else:
                if not in_table:
                    doc.add_paragraph()
            
            i += 1
        
        doc.save(word_file)
        print(f"✓ Đã tạo file Word: {word_file}")
        return True
        
    except ImportError as e:
        print(f"Lỗi: Thiếu thư viện python-docx. Cài đặt: pip install python-docx")
        return False
    except Exception as e:
        print(f"Lỗi khi chuyển sang Word: {e}")
        import traceback
        traceback.print_exc()
        return False

def convert_to_pdf_via_html(md_file, pdf_file):
    """Chuyển đổi sang PDF qua HTML"""
    try:
        from markdown import markdown
        from weasyprint import HTML, CSS
        from weasyprint.text.fonts import FontConfiguration
        
        print("Đang chuyển đổi sang PDF...")
        
        # Đọc markdown
        with open(md_file, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # Chuyển markdown sang HTML
        html_content = markdown(md_content, extensions=['tables', 'fenced_code', 'codehilite'])
        
        # HTML template với CSS
        html_template = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <style>
        @page {{
            size: A4;
            margin: 2cm;
        }}
        body {{
            font-family: "Times New Roman", serif;
            font-size: 11pt;
            line-height: 1.6;
            color: #333;
        }}
        h1 {{
            color: #2c3e50;
            border-bottom: 3px solid #3498db;
            padding-bottom: 10px;
            margin-top: 30px;
            page-break-after: avoid;
        }}
        h2 {{
            color: #34495e;
            margin-top: 25px;
            border-bottom: 2px solid #95a5a6;
            padding-bottom: 5px;
            page-break-after: avoid;
        }}
        h3 {{
            color: #7f8c8d;
            margin-top: 20px;
            page-break-after: avoid;
        }}
        h4, h5, h6 {{
            color: #95a5a6;
            margin-top: 15px;
            page-break-after: avoid;
        }}
        code {{
            background-color: #f4f4f4;
            padding: 2px 6px;
            border-radius: 3px;
            font-family: "Courier New", monospace;
            font-size: 9pt;
            color: #c7254e;
        }}
        pre {{
            background-color: #f8f8f8;
            border: 1px solid #ddd;
            border-radius: 5px;
            padding: 10px;
            overflow-x: auto;
            page-break-inside: avoid;
        }}
        pre code {{
            background-color: transparent;
            padding: 0;
            color: #333;
        }}
        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 20px 0;
            page-break-inside: avoid;
        }}
        th, td {{
            border: 1px solid #ddd;
            padding: 8px 12px;
            text-align: left;
        }}
        th {{
            background-color: #3498db;
            color: white;
            font-weight: bold;
        }}
        tr:nth-child(even) {{
            background-color: #f9f9f9;
        }}
        ul, ol {{
            margin-left: 20px;
        }}
        hr {{
            border: none;
            border-top: 1px solid #ddd;
            margin: 20px 0;
        }}
        p {{
            margin: 10px 0;
            text-align: justify;
        }}
    </style>
</head>
<body>
{html_content}
</body>
</html>"""
        
        HTML(string=html_template).write_pdf(pdf_file)
        print(f"✓ Đã tạo file PDF: {pdf_file}")
        return True
        
    except ImportError as e:
        print(f"⚠ Không thể tạo PDF: Thiếu thư viện. Cài đặt: pip install weasyprint markdown")
        return False
    except Exception as e:
        print(f"Lỗi khi chuyển sang PDF: {e}")
        import traceback
        traceback.print_exc()
        return False

def main():
    md_file = '/home/tth193/Documents/h5_code/00_CODE_H5/iq_tcp/BAO_CAO_IQTCP_H5.md'
    
    if not os.path.exists(md_file):
        print(f"Lỗi: File không tồn tại: {md_file}")
        sys.exit(1)
    
    base_name = os.path.splitext(md_file)[0]
    word_file = f"{base_name}.docx"
    pdf_file = f"{base_name}.pdf"
    
    print("=" * 80)
    print("CHUYỂN ĐỔI FILE MARKDOWN SANG WORD VÀ PDF")
    print("=" * 80)
    print(f"File nguồn: {md_file}\n")
    
    # Thử pypandoc trước
    if not convert_with_pypandoc(md_file, word_file, pdf_file):
        # Nếu không có pypandoc, dùng python-docx
        print("\n" + "-" * 80)
        convert_with_python_docx(md_file, word_file)
        
        # Chuyển sang PDF
        print("\n" + "-" * 80)
        convert_to_pdf_via_html(md_file, pdf_file)
    
    print("\n" + "=" * 80)
    print("HOÀN THÀNH")
    print("=" * 80)
    
    if os.path.exists(word_file):
        print(f"✓ File Word: {word_file}")
    if os.path.exists(pdf_file):
        print(f"✓ File PDF: {pdf_file}")

if __name__ == '__main__':
    main()

